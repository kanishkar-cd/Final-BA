"""Word document export functionality."""

from __future__ import annotations

import logging
from datetime import datetime
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

from export.formatter import StoryFormatter
from export.models import ExportRequest, ExportResponse, ExportStatus, StoryExportData
from export.utils import ensure_output_directory, generate_export_filename, get_output_base_dir

logger = logging.getLogger("export.word_export")


class WordExporter:
    """Exports user stories to Word (.docx) format with comprehensive details.
    
    Generates documents containing:
    - Epic and Feature information
    - User Stories with full details
    - Acceptance Criteria
    - Business Rules
    - Dependencies
    - Priority and Story Points
    - Additional metadata (persona, business value, confidence, etc.)
    """

    def __init__(self) -> None:
        self.base_dir = get_output_base_dir()
        self.output_dir = ensure_output_directory(self.base_dir, "word")

    def export(self, request: ExportRequest) -> ExportResponse:
        """Export user stories to a Word document.

        Parameters
        ----------
        request:
            Export request with stories and options.

        Returns
        -------
        ExportResponse
            Export result with file path or error.
        """
        export_id = f"word_{datetime.utcnow().strftime('%Y%m%d%H%M%S')}"
        logger.info("Starting Word export: %s", export_id)

        try:
            filename = request.output_filename or generate_export_filename(
                request.project_name, "docx"
            )
            file_path = self.output_dir / filename

            doc = self._create_document(request)
            doc.save(str(file_path))

            logger.info("Word export completed: %s", file_path)
            return ExportResponse(
                export_id=export_id,
                status=ExportStatus.COMPLETED,
                format=request.format,
                file_path=str(file_path),
                story_count=len(request.stories),
                completed_at=datetime.utcnow(),
            )

        except Exception as exc:
            logger.exception("Word export failed: %s", exc)
            return ExportResponse(
                export_id=export_id,
                status=ExportStatus.FAILED,
                format=request.format,
                error_message=str(exc),
                story_count=len(request.stories),
            )

    def _create_document(self, request: ExportRequest) -> Document:
        """Create a Word document with formatted user stories.

        Parameters
        ----------
        request:
            Export request.

        Returns
        -------
        Document
            Formatted Word document.
        """
        doc = Document()

        # Add title
        title = doc.add_heading(f"{request.project_name}", level=0)
        title_run = title.runs[0]
        title_run.font.color.rgb = RGBColor(31, 73, 125)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        subtitle = doc.add_heading("User Stories Export", level=1)
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Add generation timestamp
        timestamp_para = doc.add_paragraph()
        timestamp_run = timestamp_para.add_run(
            f"Generated: {datetime.utcnow().strftime('%Y-%m-%d %H:%M:%S UTC')}"
        )
        timestamp_run.italic = True
        timestamp_run.font.size = Pt(10)
        timestamp_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Add summary
        summary_para = doc.add_paragraph()
        summary_para.add_run(f"Total Stories: {len(request.stories)}").bold = True
        summary_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc.add_paragraph()  # Spacer

        # Add table of contents placeholder
        toc_heading = doc.add_heading("Table of Contents", level=1)
        toc_para = doc.add_paragraph()
        toc_para.add_run("User stories are organized sequentially below.").italic = True

        doc.add_page_break()

        # Add each story
        for idx, story in enumerate(request.stories, start=1):
            if idx > 1:
                doc.add_page_break()
            self._add_story_to_document(doc, story, idx, request.include_metadata)

        return doc

    def _add_story_to_document(
        self, doc: Document, story: StoryExportData, story_number: int, include_metadata: bool
    ) -> None:
        """Add a single user story to the document with comprehensive details.

        Parameters
        ----------
        doc:
            Word document.
        story:
            User story data.
        story_number:
            Sequential story number.
        include_metadata:
            Whether to include metadata section.
        """
        # Story title with number
        title_text = f"{story_number}. {StoryFormatter.format_title(story, include_id=True)}"
        title = doc.add_heading(title_text, level=1)
        title_run = title.runs[0]
        title_run.font.color.rgb = RGBColor(46, 116, 181)

        # Epic and Feature (if available)
        if story.epic or story.feature:
            epic_feature_para = doc.add_paragraph()
            if story.epic:
                epic_run = epic_feature_para.add_run(f"📋 Epic: ")
                epic_run.bold = True
                epic_feature_para.add_run(f"{story.epic}    ")
            if story.feature:
                feature_run = epic_feature_para.add_run(f"🎯 Feature: ")
                feature_run.bold = True
                epic_feature_para.add_run(story.feature)

        # Priority and Story Points table
        self._add_priority_points_table(doc, story)

        doc.add_paragraph()  # Spacer

        if story.user_story:
            us_heading = doc.add_heading("📖 User Story", level=2)
            us_para = doc.add_paragraph(story.user_story)
            us_para.paragraph_format.line_spacing = 1.15

        # Description section
        desc_heading = doc.add_heading("📝 Description", level=2)
        desc_para = doc.add_paragraph(story.description)
        desc_para.paragraph_format.line_spacing = 1.15

        # Acceptance Criteria section
        self._add_acceptance_criteria_section(doc, story)

        # Business Rules section
        if story.business_rules:
            self._add_business_rules_section(doc, story.business_rules)

        # Dependencies section
        if story.dependencies:
            self._add_dependencies_section(doc, story.dependencies)

        # Additional Sections
        self._add_additional_sections(doc, story)

        # Metadata section
        if include_metadata:
            self._add_metadata_section(doc, story)

        # Separator
        doc.add_paragraph()

    def _add_priority_points_table(self, doc: Document, story: StoryExportData) -> None:
        """Add a table with priority and story points.

        Parameters
        ----------
        doc:
            Word document.
        story:
            User story data.
        """
        table = doc.add_table(rows=1, cols=4)
        table.style = "Light Grid Accent 1"

        # Header row
        header_cells = table.rows[0].cells
        header_cells[0].text = "Story ID"
        header_cells[1].text = "Priority"
        header_cells[2].text = "Story Points"
        header_cells[3].text = "Status"

        # Make header bold
        for cell in header_cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True

        # Data row
        row_cells = table.add_row().cells
        row_cells[0].text = story.story_id
        row_cells[1].text = story.priority or "Not Set"
        row_cells[2].text = str(story.story_points) if story.story_points else "Not Set"
        
        # Status based on labels or metadata
        status = "Active"
        if "has_risks" in story.labels:
            status = "⚠️ Has Risks"
        row_cells[3].text = status

    def _add_acceptance_criteria_section(self, doc: Document, story: StoryExportData) -> None:
        """Add acceptance criteria section.

        Parameters
        ----------
        doc:
            Word document.
        story:
            User story data.
        """
        ac_heading = doc.add_heading("✅ Acceptance Criteria", level=2)
        
        if story.acceptance_criteria:
            for idx, criterion in enumerate(story.acceptance_criteria, start=1):
                para = doc.add_paragraph(criterion, style="List Number")
                para.paragraph_format.left_indent = Inches(0.25)
        else:
            para = doc.add_paragraph("No acceptance criteria defined.")
            para.italic = True

    def _add_business_rules_section(self, doc: Document, business_rules: list[str]) -> None:
        """Add business rules section.

        Parameters
        ----------
        doc:
            Word document.
        business_rules:
            List of business rules.
        """
        if not business_rules:
            return

        br_heading = doc.add_heading("📜 Business Rules", level=2)
        
        for idx, rule in enumerate(business_rules, start=1):
            para = doc.add_paragraph(rule, style="List Bullet")
            para.paragraph_format.left_indent = Inches(0.25)

    def _add_dependencies_section(self, doc: Document, dependencies: list[Any]) -> None:
        """Add dependencies section.

        Parameters
        ----------
        doc:
            Word document.
        dependencies:
            List of dependencies (can be dicts, strings, or objects).
        """
        if not dependencies:
            return

        dep_heading = doc.add_heading("🔗 Dependencies", level=2)
        
        # Create table for dependencies
        table = doc.add_table(rows=1, cols=3)
        table.style = "Light List Accent 1"

        # Header row
        header_cells = table.rows[0].cells
        header_cells[0].text = "Dependency ID"
        header_cells[1].text = "Type"
        header_cells[2].text = "Description"

        # Make header bold
        for cell in header_cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True

        # Add dependency rows
        for dep in dependencies:
            row_cells = table.add_row().cells
            
            if isinstance(dep, dict):
                row_cells[0].text = dep.get("id", "N/A")
                row_cells[1].text = dep.get("type", "unknown")
                row_cells[2].text = dep.get("description", "")
            elif isinstance(dep, str):
                row_cells[0].text = dep
                row_cells[1].text = "reference"
                row_cells[2].text = ""
            else:
                row_cells[0].text = str(dep)
                row_cells[1].text = "unknown"
                row_cells[2].text = ""

    def _add_additional_sections(self, doc: Document, story: StoryExportData) -> None:
        """Add additional sections like risks, assumptions, definition of done.

        Parameters
        ----------
        doc:
            Word document.
        story:
            User story data.
        """
        # Risks
        risks = story.metadata.get("risks", [])
        if risks:
            risk_heading = doc.add_heading("⚠️ Risks", level=2)
            for risk in risks:
                para = doc.add_paragraph(risk, style="List Bullet")
                para.paragraph_format.left_indent = Inches(0.25)

        # Assumptions
        assumptions = story.metadata.get("assumptions", [])
        if assumptions:
            assump_heading = doc.add_heading("💭 Assumptions", level=2)
            for assumption in assumptions:
                para = doc.add_paragraph(assumption, style="List Bullet")
                para.paragraph_format.left_indent = Inches(0.25)

        # Definition of Done
        if story.definition_of_done:
            dod_heading = doc.add_heading("✔️ Definition of Done", level=2)
            for item in story.definition_of_done:
                para = doc.add_paragraph(item, style="List Bullet")
                para.paragraph_format.left_indent = Inches(0.25)

    def _add_metadata_section(self, doc: Document, story: StoryExportData) -> None:
        """Add comprehensive metadata section.

        Parameters
        ----------
        doc:
            Word document.
        story:
            User story data.
        """
        # Build metadata rows
        metadata_rows = []

        # Basic metadata
        if story.assignee:
            metadata_rows.append(("Assignee", story.assignee))
        
        if story.created_at:
            metadata_rows.append(("Created At", story.created_at.strftime("%Y-%m-%d %H:%M:%S")))

        if story.labels:
            metadata_rows.append(("Labels", ", ".join(story.labels)))

        # Extended metadata from story.metadata
        if "confidence_score" in story.metadata:
            score = story.metadata["confidence_score"]
            metadata_rows.append(("Confidence Score", f"{score:.2%}"))

        if "business_value" in story.metadata:
            metadata_rows.append(("Business Value", story.metadata["business_value"]))

        if "persona" in story.metadata:
            metadata_rows.append(("Persona", story.metadata["persona"]))

        if "goal" in story.metadata:
            metadata_rows.append(("Goal", story.metadata["goal"]))

        if "chunk_ids_used" in story.metadata:
            chunk_ids = story.metadata["chunk_ids_used"]
            if chunk_ids:
                metadata_rows.append(("Source Chunks", f"{len(chunk_ids)} chunks"))

        if story.traceability:
            trace = story.traceability
            if isinstance(trace, dict):
                if "requirements" in trace:
                    reqs = trace["requirements"]
                    if isinstance(reqs, list):
                        metadata_rows.append(("Requirements", ", ".join(str(r) for r in reqs)))

        # Only add section if we have metadata to show
        if metadata_rows:
            meta_heading = doc.add_heading("ℹ️ Metadata", level=2)
            
            table = doc.add_table(rows=len(metadata_rows), cols=2)
            table.style = "Light Grid"

            for idx, (key, value) in enumerate(metadata_rows):
                row = table.rows[idx]
                key_cell = row.cells[0]
                value_cell = row.cells[1]
                
                # Key cell (bold)
                key_para = key_cell.paragraphs[0]
                key_run = key_para.add_run(key)
                key_run.font.bold = True
                
                # Value cell
                value_cell.text = str(value)
